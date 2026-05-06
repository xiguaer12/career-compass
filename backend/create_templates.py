"""Generate real, usable career planning templates as .docx and .xlsx files."""
import sys
import os

# Fix Windows GBK encoding issue
if sys.platform == "win32":
    import io
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill, numbers
from openpyxl.utils import get_column_letter

OUT = os.path.join(os.path.dirname(__file__), "src", "main", "resources", "static", "templates")
os.makedirs(OUT, exist_ok=True)

thin = Side(style="thin")
border_all = Border(left=thin, right=thin, top=thin, bottom=thin)
header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
header_font = Font(name="微软雅黑", bold=True, color="FFFFFF", size=11)
cell_font = Font(name="微软雅黑", size=10)
title_font_xlsx = Font(name="微软雅黑", bold=True, size=14)

def style_header(ws, row, cols):
    for c in range(1, cols + 1):
        cell = ws.cell(row=row, column=c)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        cell.border = border_all

def style_data(ws, start_row, end_row, cols):
    for r in range(start_row, end_row + 1):
        for c in range(1, cols + 1):
            cell = ws.cell(row=r, column=c)
            cell.font = cell_font
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            cell.border = border_all

def set_docx_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(6)

def add_docx_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(18)

def add_docx_section(doc, title):
    h = doc.add_heading(title, level=2)
    for run in h.runs:
        run.font.size = Pt(14)

# ═══════════════════════════════════════════════════════════════
# 1. 通用校招简历模板 (employment-resume-template.docx)
# ═══════════════════════════════════════════════════════════════
doc = Document()
set_docx_style(doc)
add_docx_title(doc, "通用校招简历模板")
doc.add_paragraph("使用说明：请将下方占位内容替换为您的真实信息，删除本行说明后保存为 PDF 投递。", style="Subtitle")

add_docx_section(doc, "一、基本信息")
table = doc.add_table(rows=5, cols=4, style="Light Grid Accent 1")
info = [
    ("姓名：__________", "性别：__________", "出生年月：__________", "政治面貌：__________"),
    ("手机：__________", "邮箱：__________", "籍贯：__________", "现居城市：__________"),
    ("学历：__________", "专业：__________", "毕业院校：__________", "毕业时间：__________"),
    ("GPA：__________", "排名：__________", "英语水平：__________", "计算机水平：__________"),
]
for i, row_data in enumerate(info):
    for j, text in enumerate(row_data):
        table.cell(i, j).text = text

add_docx_section(doc, "二、求职意向")
table2 = doc.add_table(rows=4, cols=4, style="Light Grid Accent 1")
intent = [
    ("意向岗位：__________", "意向城市：__________", "期望薪资：__________", "到岗时间：__________"),
    ("意向行业：__________", "工作性质：全职/实习", "", ""),
]
for i, row_data in enumerate(intent):
    for j, text in enumerate(row_data):
        if text:
            table2.cell(i, j).text = text

add_docx_section(doc, "三、教育经历")
doc.add_paragraph("__________大学  __________专业  本科/硕士  20__年9月 - 20__年6月")
doc.add_paragraph("主修课程：________________________________________")
doc.add_paragraph("荣誉奖项：________________________________________", style="List Bullet")
doc.add_paragraph("________________________________________________", style="List Bullet")

add_docx_section(doc, "四、实习/项目经历")
for i in range(1, 3):
    doc.add_paragraph(f"经历 {i}：__________公司/机构  __________岗位  20__年__月 - 20__年__月")
    doc.add_paragraph("工作内容：", style="List Bullet")
    doc.add_paragraph("1. ________________________________________________")
    doc.add_paragraph("2. ________________________________________________")
    doc.add_paragraph("3. ________________________________________________")

add_docx_section(doc, "五、校园经历")
doc.add_paragraph("__________社团/组织  __________职务  20__年__月 - 20__年__月")
doc.add_paragraph("主要成果：________________________________________")

add_docx_section(doc, "六、技能与证书")
doc.add_paragraph("语言能力：________________________________________")
doc.add_paragraph("专业技能：________________________________________")
doc.add_paragraph("证书：____________________________________________")

add_docx_section(doc, "七、自我评价")
doc.add_paragraph("（建议控制在 3-5 行，突出与意向岗位的匹配点）")

doc.save(os.path.join(OUT, "employment-resume-template.docx"))
print("✓ employment-resume-template.docx")

# ═══════════════════════════════════════════════════════════════
# 2. 考公报名信息整理表 (civil-position-screening.xlsx)
# ═══════════════════════════════════════════════════════════════
wb = Workbook()
ws = wb.active
ws.title = "岗位筛选表"

title_row = ["序号", "招录单位", "职位名称", "职位代码", "招录人数", "专业要求", "学历要求", "政治面貌",
             "基层工作年限", "其他条件", "报名时间", "笔试时间", "面试时间", "备注", "适合度评级"]
for c, val in enumerate(title_row, 1):
    ws.cell(row=1, column=c, value=val)
style_header(ws, 1, len(title_row))

for r in range(2, 12):
    ws.cell(row=r, column=1, value=r - 1)
style_data(ws, 2, 11, len(title_row))

col_widths = [5, 18, 16, 14, 8, 16, 10, 10, 12, 18, 14, 14, 14, 18, 10]
for i, w in enumerate(col_widths, 1):
    ws.column_dimensions[get_column_letter(i)].width = w
ws.sheet_properties.tabColor = "2563eb"
wb.save(os.path.join(OUT, "civil-position-screening.xlsx"))
print("✓ civil-position-screening.xlsx")

# ═══════════════════════════════════════════════════════════════
# 3. 复试个人陈述模板 (postgraduate-personal-statement.docx)
# ═══════════════════════════════════════════════════════════════
doc = Document()
set_docx_style(doc)
add_docx_title(doc, "硕士研究生复试个人陈述")
doc.add_paragraph("报考院校：____________  报考专业：____________  考生编号：____________")

add_docx_section(doc, "一、个人基本情况")
doc.add_paragraph("姓名：____________  性别：______  出生年月：____________")
doc.add_paragraph("本科院校：____________  本科专业：____________  预计毕业时间：____________")
doc.add_paragraph("政治面貌：____________  联系方式：____________")

add_docx_section(doc, "二、本科阶段学习情况")
doc.add_paragraph("GPA：______  专业排名：______/______")
doc.add_paragraph("主修课程：")
doc.add_paragraph("（列举与报考专业密切相关的 8-12 门课程及成绩）")
for i in range(1, 9):
    doc.add_paragraph(f"  {i}. 《____________》  成绩：______", style="List Bullet")
doc.add_paragraph("外语水平：CET-4 ______ 分  CET-6 ______ 分  其他：____________")
doc.add_paragraph("计算机水平：____________")

add_docx_section(doc, "三、科研与竞赛经历")
doc.add_paragraph("1. 项目名称：________________________________________")
doc.add_paragraph("   角色：__________  时间：20__年__月 - 20__年__月")
doc.add_paragraph("   成果：________________________________________")
doc.add_paragraph("2. 竞赛/论文/专利：__________________________________")
doc.add_paragraph("   获奖情况：________________________________________")

add_docx_section(doc, "四、社会实践与实习")
doc.add_paragraph("1. __________单位  __________岗位  20__年__月 - 20__年__月")
doc.add_paragraph("   主要工作：________________________________________")
doc.add_paragraph("2. ________________________________________________")

add_docx_section(doc, "五、研究生阶段学习计划")
doc.add_paragraph("研究方向：________________________________________")
doc.add_paragraph("研一目标：________________________________________")
doc.add_paragraph("研二目标：________________________________________")
doc.add_paragraph("研三目标：________________________________________")

add_docx_section(doc, "六、个人特点与优势")
doc.add_paragraph("（简述 2-3 个核心优势，每个附简短例证）")

add_docx_section(doc, "七、其他补充")
doc.add_paragraph("（如有需要说明的情况，如跨专业原因、GAP 经历等）")

doc.save(os.path.join(OUT, "postgraduate-personal-statement.docx"))
print("✓ postgraduate-personal-statement.docx")

# ═══════════════════════════════════════════════════════════════
# 4. 面试复盘记录表 (employment-interview-review.docx)
# ═══════════════════════════════════════════════════════════════
doc = Document()
set_docx_style(doc)
add_docx_title(doc, "面试复盘记录表")
doc.add_paragraph("坚持每次面试后复盘，持续优化面试表现", style="Subtitle")

add_docx_section(doc, "面试基本信息")
table = doc.add_table(rows=4, cols=4, style="Light Grid Accent 1")
items = [
    ("公司名称：__________", "岗位：__________", "面试轮次：第____轮", "面试日期：20__年__月__日"),
    ("面试形式：□现场 □视频 □电话", "面试官角色：__________", "面试时长：____分钟", "是否笔试：□是 □否"),
]
for i, row_data in enumerate(items):
    for j, text in enumerate(row_data):
        table.cell(i, j).text = text

add_docx_section(doc, "面试前准备")
doc.add_paragraph("对公司了解程度：□充分 □一般 □不足", style="List Bullet")
doc.add_paragraph("岗位职责理解：□清晰 □模糊", style="List Bullet")
doc.add_paragraph("准备的 3 个关键点：", style="List Bullet")
doc.add_paragraph("  1. ________________________________________________")
doc.add_paragraph("  2. ________________________________________________")
doc.add_paragraph("  3. ________________________________________________")

add_docx_section(doc, "面试中被问到的问题（逐题记录）")
for i in range(1, 9):
    doc.add_paragraph(f"Q{i}：________________________________________________")
    doc.add_paragraph(f"A{i} 要点：________________________________________")
    doc.add_paragraph(f"自评：□答得好  □一般  □答得不好    改进：________________", style="List Bullet")

add_docx_section(doc, "我的提问")
for i in range(1, 4):
    doc.add_paragraph(f"{i}. ________________________________________________")

add_docx_section(doc, "整体复盘")
doc.add_paragraph("表现好的地方：________________________________________")
doc.add_paragraph("需要改进的地方：________________________________________")
doc.add_paragraph("下次面试要特别注意的 3 件事：")
doc.add_paragraph("  1. ________________________________________________")
doc.add_paragraph("  2. ________________________________________________")
doc.add_paragraph("  3. ________________________________________________")
doc.add_paragraph(f"综合自评分：____/10    是否有后续：□是 □否    预计通知时间：20__年__月__日")

doc.save(os.path.join(OUT, "employment-interview-review.docx"))
print("✓ employment-interview-review.docx")

# ═══════════════════════════════════════════════════════════════
# 5. 校招投递跟踪表 (employment-application-tracker.xlsx)
# ═══════════════════════════════════════════════════════════════
wb = Workbook()
ws = wb.active
ws.title = "校招投递跟踪"

title_row = ["序号", "公司名称", "投递岗位", "投递时间", "投递渠道", "当前状态",
             "笔试时间", "面试时间", "Offer情况", "备注"]
for c, val in enumerate(title_row, 1):
    ws.cell(row=1, column=c, value=val)
style_header(ws, 1, len(title_row))

statuses = ["简历投递", "筛选通过", "笔试完成", "一面", "二面", "三面", "HR面", "Offer", "已拒绝", "已接受"]
for r in range(2, 32):
    ws.cell(row=r, column=1, value=r - 1)
style_data(ws, 2, 31, len(title_row))

col_widths = [5, 18, 16, 14, 12, 12, 14, 14, 12, 20]
for i, w in enumerate(col_widths, 1):
    ws.column_dimensions[get_column_letter(i)].width = w

# Add a summary section
ws_summary = wb.create_sheet("统计")
ws_summary.cell(row=1, column=1, value="校招投递统计").font = title_font_xlsx
summary_headers = ["指标", "数量"]
for c, val in enumerate(summary_headers, 1):
    ws_summary.cell(row=3, column=c, value=val)
style_header(ws_summary, 3, 2)
summary_data = [["投递总数", "0"], ["筛选通过", "0"], ["笔试", "0"], ["面试", "0"], ["Offer", "0"], ["已拒绝", "0"]]
for i, (label, val) in enumerate(summary_data):
    ws_summary.cell(row=4 + i, column=1, value=label)
    ws_summary.cell(row=4 + i, column=2, value=val)
style_data(ws_summary, 4, 9, 2)
ws_summary.column_dimensions["A"].width = 14
ws_summary.column_dimensions["B"].width = 10

wb.save(os.path.join(OUT, "employment-application-tracker.xlsx"))
print("✓ employment-application-tracker.xlsx")

# ═══════════════════════════════════════════════════════════════
# 6. 考公备考周计划模板 (civil-study-plan.docx)
# ═══════════════════════════════════════════════════════════════
doc = Document()
set_docx_style(doc)
add_docx_title(doc, "考公备考周计划模板")
doc.add_paragraph("备考目标考试：□国考  □省考（______省）  □选调生", style="Subtitle")

add_docx_section(doc, "总体规划")
doc.add_paragraph("备考总周期：20__年__月 - 20__年__月（共____周）")
doc.add_paragraph("每天学习时长：工作日 ____小时 / 周末 ____小时")
doc.add_paragraph("薄弱模块（需重点攻克）：____________________________")

add_docx_section(doc, "每日时间表示例")
timetable = [
    ("06:30 - 07:00", "晨读（时政、常识）"),
    ("07:00 - 08:00", "早餐 + 通勤"),
    ("08:00 - 11:30", "行测专项训练"),  # corrected
    ("11:30 - 13:00", "午餐 + 休息"),
    ("13:00 - 15:00", "申论专项训练"),  # corrected
    ("15:00 - 17:30", "真题/模拟题"),
    ("17:30 - 19:00", "晚餐 + 运动"),
    ("19:00 - 21:00", "错题回顾 / 查漏补缺"),
    ("21:00 - 22:00", "复盘当日学习 + 列明日计划"),
]
table = doc.add_table(rows=len(timetable) + 1, cols=2, style="Light Grid Accent 1")
table.cell(0, 0).text = "时间"
table.cell(0, 1).text = "内容"
for i, (t, c) in enumerate(timetable):
    table.cell(i + 1, 0).text = t
    table.cell(i + 1, 1).text = c

add_docx_section(doc, "每周学习计划（可复制本页填写每周计划）")
doc.add_paragraph("第 ____ 周（20__年__月__日 - 20__年__月__日）")
doc.add_paragraph("本周目标：________________________________________")
plan_table = [
    ("周一", "行测：__________  申论：__________  题量：______"),
    ("周二", "行测：__________  申论：__________  题量：______"),
    ("周三", "行测：__________  申论：__________  题量：______"),
    ("周四", "行测：__________  申论：__________  题量：______"),
    ("周五", "行测：__________  申论：__________  题量：______"),
    ("周六", "行测：__________  申论：__________  题量：______"),
    ("周日", "本周复盘 + 错题整理 + 下周计划"),
]
t2 = doc.add_table(rows=len(plan_table) + 1, cols=2, style="Light Grid Accent 1")
t2.cell(0, 0).text = "日期"
t2.cell(0, 1).text = "计划内容"
for i, (day, plan) in enumerate(plan_table):
    t2.cell(i + 1, 0).text = day
    t2.cell(i + 1, 1).text = plan

add_docx_section(doc, "各模块进度跟踪")
modules = [
    ("言语理解", "____/40"), ("数量关系", "____/15"), ("判断推理", "____/40"),
    ("资料分析", "____/20"), ("常识判断", "____/20"), ("申论小题", "____"),
    ("申论大作文", "____"), ("面试表达", "____"),
]
t3 = doc.add_table(rows=len(modules) + 1, cols=3, style="Light Grid Accent 1")
t3.cell(0, 0).text = "模块"
t3.cell(0, 1).text = "目标正确率"
t3.cell(0, 2).text = "当前正确率"
for i, (mod, target) in enumerate(modules):
    t3.cell(i + 1, 0).text = mod
    t3.cell(i + 1, 1).text = target
    t3.cell(i + 1, 2).text = "____"

doc.save(os.path.join(OUT, "civil-study-plan.docx"))
print("✓ civil-study-plan.docx")

# ═══════════════════════════════════════════════════════════════
# 7. 考公报名材料核对清单 (civil-application-checklist.docx)
# ═══════════════════════════════════════════════════════════════
doc = Document()
set_docx_style(doc)
add_docx_title(doc, "考公报名材料核对清单")
doc.add_paragraph("报名前逐项核对，确保材料齐全、信息准确", style="Subtitle")

items = [
    ("一、身份证明类", [
        ("身份证原件及复印件（正反面）", "□"),
        ("户口本原件及复印件（户主页+本人页）", "□"),
        ("近期免冠证件照（电子版+纸质版，按公告要求底色尺寸）", "□"),
        ("学生证原件及复印件（应届生）", "□"),
    ]),
    ("二、学历学位类", [
        ("本科毕业证书原件及复印件", "□"),
        ("学士学位证书原件及复印件", "□"),
        ("研究生毕业证书原件及复印件（如有）", "□"),
        ("硕士学位证书原件及复印件（如有）", "□"),
        ("教育部学历证书电子注册备案表（学信网下载）", "□"),
        ("教育部学位认证报告（如需要）", "□"),
    ]),
    ("三、报考资格类", [
        ("英语四六级成绩单原件及复印件", "□"),
        ("计算机等级证书原件及复印件", "□"),
        ("职业资格证书原件及复印件（如法律职业资格等）", "□"),
        ("党员证明（所在党组织开具）", "□"),
        ("基层工作经历证明（如需要）", "□"),
        ("应届毕业生推荐表", "□"),
        ("教育部学籍在线验证报告（应届生）", "□"),
    ]),
    ("四、报名信息核对", [
        ("姓名、性别、出生日期与身份证一致", "□"),
        ("学历、学位信息与证书一致", "□"),
        ("专业名称与学信网一致（注意专业目录口径）", "□"),
        ("工作经历起止时间准确无遗漏", "□"),
        ("家庭成员信息完整准确", "□"),
        ("联系方式（手机、邮箱）正确", "□"),
    ]),
    ("五、网上报名确认", [
        ("报名信息已提交并确认", "□"),
        ("资格审核状态：□通过  □未审核  □退回修改", "□"),
        ("报名费已缴纳", "□"),
        ("报名登记表已下载保存（PDF）", "□"),
        ("报名确认页已打印（如需现场确认）", "□"),
    ]),
    ("六、考试准备", [
        ("准考证已打印（建议打印 2-3 份）", "□"),
        ("身份证原件随身携带", "□"),
        ("考场地址、交通路线已确认", "□"),
        ("考试用品准备（2B 铅笔、黑色签字笔、橡皮等）", "□"),
    ]),
]

for section_title, section_items in items:
    add_docx_section(doc, section_title)
    for item_text, checkbox in section_items:
        p = doc.add_paragraph()
        p.text = f"{checkbox}  {item_text}"

doc.add_paragraph("")
doc.add_paragraph("温馨提示：建议提前 2 周准备以上材料，避免报名截止前忙乱遗漏。")
doc.save(os.path.join(OUT, "civil-application-checklist.docx"))
print("✓ civil-application-checklist.docx")

# ═══════════════════════════════════════════════════════════════
# 8. 考研院校专业对比表 (postgraduate-school-comparison.xlsx)
# ═══════════════════════════════════════════════════════════════
wb = Workbook()
ws = wb.active
ws.title = "院校对比表"

title_row = ["对比维度", "院校A：__________", "院校B：__________", "院校C：__________"]
for c, val in enumerate(title_row, 1):
    ws.cell(row=1, column=c, value=val)
style_header(ws, 1, len(title_row))

dimensions = [
    "院校层次（985/211/双一流/普通）", "所在城市", "专业排名/学科评估等级",
    "研究方向匹配度", "导师团队实力", "招生人数（含推免）",
    "报录比（近3年）", "复试分数线（近3年）", "复试比例（差额比）",
    "初试科目", "专业课参考书目数量", "是否保护一志愿",
    "学费（年）", "奖学金覆盖比例", "住宿条件",
    "学制（年）", "毕业去向（就业/读博）", "实习/项目机会",
    "校园环境/生活成本", "综合评分（1-10）",
]
for r, dim in enumerate(dimensions, 2):
    ws.cell(row=r, column=1, value=dim)
style_data(ws, 2, len(dimensions) + 1, len(title_row))

# Highlight section headers
section_fill = PatternFill(start_color="D6E4F0", end_color="D6E4F0", fill_type="solid")
section_rows = []  # No specific section rows needed

col_widths = [26, 22, 22, 22]
for i, w in enumerate(col_widths, 1):
    ws.column_dimensions[get_column_letter(i)].width = w

# Add a scoring sheet
ws2 = wb.create_sheet("综合评分表")
ws2.cell(row=1, column=1, value="考研院校综合评分表").font = title_font_xlsx
score_headers = ["评分项（权重）", "院校A得分", "院校B得分", "院校C得分"]
for c, val in enumerate(score_headers, 1):
    ws2.cell(row=3, column=c, value=val)
style_header(ws2, 3, 4)

score_items = [
    "院校实力 (20%)", "专业水平 (20%)", "导师资源 (15%)",
    "地理位置 (10%)", "录取难度 (15%)", "费用 (10%)",
    "就业前景 (10%)", "加权总分",
]
for r, item in enumerate(score_items, 4):
    ws2.cell(row=r, column=1, value=item)
style_data(ws2, 4, 11, 4)
ws2.column_dimensions["A"].width = 22
for c in range(2, 5):
    ws2.column_dimensions[get_column_letter(c)].width = 14

wb.save(os.path.join(OUT, "postgraduate-school-comparison.xlsx"))
print("✓ postgraduate-school-comparison.xlsx")

# ═══════════════════════════════════════════════════════════════
# 9. 复试材料核对清单 (postgraduate-retest-materials.docx)
# ═══════════════════════════════════════════════════════════════
doc = Document()
set_docx_style(doc)
add_docx_title(doc, "考研复试材料核对清单")
doc.add_paragraph("报考院校：____________  报考专业：____________", style="Subtitle")

checklist = [
    ("一、身份与学历证明", [
        "身份证原件及复印件（正反面）",
        "初试准考证",
        "本科成绩单（加盖学校公章）",
        "应届生：学生证原件及复印件",
        "往届生：本科毕业证、学位证原件及复印件",
        "教育部学籍在线验证报告（应届生）",
        "教育部学历证书电子注册备案表（往届生）",
        "复试通知书/复试通知截图",
    ]),
    ("二、政审与品德", [
        "政治审查表（按要求盖章）",
        "党员证明（如报考或导师有要求）",
        "获奖证书原件及复印件（奖学金、竞赛等）",
    ]),
    ("三、学术能力证明", [
        "个人简历（复试专用，突出学术经历）",
        "个人陈述/自述（按学校要求的字数格式）",
        "本科毕业论文/设计摘要",
        "发表的论文或研究报告（如有）",
        "参与的科研项目证明（如有）",
        "专利证书（如有）",
        "英语四六级成绩单原件及复印件",
        "其他语言能力证明（雅思/托福/日语等）",
        "计算机等级证书",
        "专业相关资格证书",
    ]),
    ("四、推荐与联系", [
        "专家推荐信（通常 2 封，按学校要求）",
        "导师联系记录/邮件回复（了解导师意向）",
    ]),
    ("五、复试行程准备", [
        "复试地点、时间已确认",
        "交通/住宿已预订",
        "复试费已缴纳（如需要）",
        "体检安排已确认（如需要）",
        "调剂志愿已准备（如初试分数处于边缘）",
    ]),
    ("六、面试准备", [
        "中英文自我介绍（2-3分钟版本）",
        "研究计划/读研规划（书面+口头）",
        "专业知识复习（重点回顾本专业核心课程）",
        "常见面试问题准备（为什么选我们/为什么转专业等）",
        "着装准备（整洁得体即可，不一定正装）",
    ]),
]

for section_title, section_items in checklist:
    add_docx_section(doc, section_title)
    for item_text in section_items:
        p = doc.add_paragraph()
        p.text = f"□  {item_text}"

doc.add_paragraph("")
doc.add_paragraph("建议提前 1 个月逐项准备，复试前 3 天最终核对。祝复试顺利！")
doc.save(os.path.join(OUT, "postgraduate-retest-materials.docx"))
print("✓ postgraduate-retest-materials.docx")

print("\n✅ All 9 template files generated successfully!")
print(f"   Output directory: {OUT}")
